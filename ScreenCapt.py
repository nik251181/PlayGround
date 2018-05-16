import tkinter as Tkinter
from tkinter import constants as Tkconstants
from tkinter import filedialog as tkFileDialog
from tkinter import Label,Entry
import time
from PIL import ImageGrab
import pyHook
from docx import Document
from docx.shared import Inches
from FileValidity import is_path_exists_or_creatable
import os

direct_path=""
combined_document = Document()

def OnKeyboardEvent(event):
  if event.Key =='F6':
    takescreenshot()
  return True

def takescreenshot():
  img = ImageGrab.grab()
  moment = time.strftime("%y%m%d%H%M%S")
  saveas = os.path.join(direct_path, 'ScreenShot_' + moment + '.jpg')
  img.save(saveas)
  combined_document.add_picture(saveas, width=Inches(6))
  combined_document.add_paragraph(" ")
  combined_document.add_paragraph(" ")
  savedoc = os.path.join(direct_path,e1.get() + '.docx')
  combined_document.save(savedoc)
  os.remove(saveas)

class TkFileDialogExample(Tkinter.Frame):

  def __init__(self, root):

    Tkinter.Frame.__init__(self, root)
    button_opt = {'fill': Tkconstants.BOTH, 'padx': 5, 'pady': 5}
    Tkinter.Button(self, text='Save in!!', command=self.askdirectory).pack(**button_opt)
    self.dir_opt = options = {}
    options['initialdir'] = 'C:\\'
    options['mustexist'] = False
    options['parent'] = root
    options['title'] = 'Select location for saving screenshots'

  def askdirectory(self):
    global direct_path
    direct_path= tkFileDialog.askdirectory(**self.dir_opt)

if __name__=='__main__':
  root = Tkinter.Tk()
  root.title('Screen capture')
  root.resizable(width=0,height=0)
  root.geometry('{}x{}'.format(365,50))
  Label(root, text="Doc Name").grid(row=0)
  e1 = Entry(root,width=45)
  e1.grid(row=0, column=1)
  TkFileDialogExample(root).grid(row=1)
  if str(is_path_exists_or_creatable(str(os.path.join(direct_path,e1.get() + '.docx'))))=='False':
    Label(root,text="Please correct the file name", font=("Helvetica", 10, "bold italic")).grid(row=2)
  else:
    hm = pyHook.HookManager()
    hm.KeyDown = OnKeyboardEvent
    hm.HookKeyboard()
  root.mainloop()
